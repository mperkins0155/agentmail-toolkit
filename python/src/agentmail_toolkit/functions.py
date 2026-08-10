from typing import Any, Dict
from urllib.parse import urlparse
from urllib.request import HTTPRedirectHandler, build_opener
from agentmail import AgentMail
from agentmail.inboxes import CreateInboxRequest

import io
import logging
import filetype
import pymupdf
import docx


logger = logging.getLogger(__name__)


class _NoRedirectHandler(HTTPRedirectHandler):
    """Refuse redirects: following one could silently downgrade the https-only
    check on the attachment download URL to an arbitrary target scheme."""

    def redirect_request(self, req, fp, code, msg, headers, newurl):
        return None


# Same call signature as urllib.request.urlopen; kept under the same name so it
# stays the single patch point for tests.
urlopen = build_opener(_NoRedirectHandler).open

Kwargs = Dict[str, Any]

# Mirrors the AgentMail API's own enforced content ceiling (RESPONSE_SIZE_LIMIT in
# agentmail-api/src/agentmail/utils/limits.ts: 5.95 MB). The API inlines extracted
# content only up to this size and otherwise returns a URL; the toolkit inlines
# extracted attachment text into a tool result, so it uses the same ceiling rather
# than an invented one.
MAX_ATTACHMENT_BYTES = int(5.95 * 1024 * 1024)


def list_inboxes(client: AgentMail, kwargs: Kwargs):
    return client.inboxes.list(**kwargs)


def get_inbox(client: AgentMail, kwargs: Kwargs):
    return client.inboxes.get(**kwargs)


def create_inbox(client: AgentMail, kwargs: Kwargs):
    return client.inboxes.create(request=CreateInboxRequest(**kwargs))


def delete_inbox(client: AgentMail, kwargs: Kwargs):
    return client.inboxes.delete(**kwargs)


def list_threads(client: AgentMail, kwargs: Kwargs):
    return client.inboxes.threads.list(**kwargs)


def get_thread(client: AgentMail, kwargs: Kwargs):
    return client.inboxes.threads.get(**kwargs)


def get_attachment(client: AgentMail, kwargs: Kwargs):
    attachment = client.inboxes.threads.get_attachment(
        inbox_id=kwargs["inbox_id"],
        thread_id=kwargs["thread_id"],
        attachment_id=kwargs["attachment_id"],
    )

    if attachment.size > MAX_ATTACHMENT_BYTES:
        return attachment

    if urlparse(attachment.download_url).scheme != "https":
        return attachment

    try:
        # Read one byte past the cap so an oversized body is detected without
        # trusting Content-Length, which urlopen doesn't enforce itself.
        file_bytes = urlopen(attachment.download_url, timeout=15).read(
            MAX_ATTACHMENT_BYTES + 1
        )
        if len(file_bytes) > MAX_ATTACHMENT_BYTES:
            return attachment

        file_kind = filetype.guess(file_bytes)
        file_type = file_kind.mime if file_kind else None

        text = None
        if file_type == "application/pdf":
            text = ""
            for page in pymupdf.Document(stream=file_bytes):
                text += page.get_text() + "\n"
        if file_type == "application/zip":
            text = ""
            for paragraph in docx.Document(io.BytesIO(file_bytes)).paragraphs:
                text += paragraph.text + "\n"

        if text is not None:
            return attachment.model_copy(update={"text": text})
    except Exception as e:
        logger.warning(
            "attachment extraction failed for attachment_id=%s: %s",
            kwargs.get("attachment_id"),
            e,
        )

    return attachment


def send_message(client: AgentMail, kwargs: Kwargs):
    return client.inboxes.messages.send(**kwargs)


def reply_to_message(client: AgentMail, kwargs: Kwargs):
    kwargs = dict(kwargs)
    inbox_id = kwargs.pop("inbox_id")
    message_id = kwargs.pop("message_id")
    return client.inboxes.messages.reply(inbox_id, message_id, **kwargs)


def forward_message(client: AgentMail, kwargs: Kwargs):
    kwargs = dict(kwargs)
    inbox_id = kwargs.pop("inbox_id")
    message_id = kwargs.pop("message_id")
    return client.inboxes.messages.forward(inbox_id, message_id, **kwargs)


def update_message(client: AgentMail, kwargs: Kwargs):
    return client.inboxes.messages.update(**kwargs)


def agent_verify(client: AgentMail, kwargs: Kwargs):
    return client.agent.verify(**kwargs)
